
import os
import pandas as pd                 #For reading CSV and excel files
from pypdf import PdfReader         #For reading PDF files
from docx import Document   
from src.chunking import chunk_text
from src.constants import CHUNK_SIZE


#--------------------FILE READER----------------------------
def read_txt(filepath):
    '''Reads a plain text file and chunks it smartly
    1.WHAT :- read the .txt file and chunks it smartly
    2.RETURNS :- list of (chunks_text, metadata_dict)
    '''

    filename = os.path.basename(filepath)
    print(f"Reading text:{filename}")
    
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    
    #first trying to split by double newline (paragraph)
    paragraphs = [p.strip() for p in content .split("\n\n") if p.strip()]
    
    results = []
    chunk_index = 0
    
    for para_index, paragraph in enumerate(paragraphs):
        #if paragraph is small - keep in one chunk
        if len(paragraph) <= CHUNK_SIZE:
            metadata = {
                "source" : filename,
                "file_type": "txt",
                "paragraph": para_index,
                "chunk_index": chunk_index,
                "category" : "text"
            }
            results.append((paragraph,metadata))
            chunk_index += 1
        else:
            # Paragraph too long — chunk further
            sub_chunks = chunk_text(paragraph)
            for sub_index, sub_chunk in enumerate(sub_chunks):
                metadata = {
                    "source": filename,
                    "file_type": "txt",
                    "paragraph": para_index,
                    "sub_chunk": sub_index,
                    "chunk_index": chunk_index,
                    "category": "text"
                }
                results.append((sub_chunk, metadata))
                chunk_index += 1
    
    print(f"   -> {len(results)}chunks created")
    return results

#----------------------read pdf----------------------------
def read_pdf(filepath):
    filename = os.path.basename(filepath)
    print(f"reading pdf: {filename}")
    
    reader = PdfReader(filepath)
    results = []
    chunk_index = 0
    
    for page_num, page in enumerate(reader.pages):
        #extracting text from the page
        page_text = page.extract_text()
        
        if not page_text or not page_text.strip():
            continue
        
        chunks = chunk_text(page_text)
        for sub_index, chunk in enumerate(chunks):
            metadata = {
                "source": filename,
                "file_type": "pdf",
                "page": page_num + 1,          # page numbers start at 1
                "sub_chunk": sub_index,
                "chunk_index": chunk_index,
                "category": "document"
            }
            results.append((chunk, metadata))
            chunk_index += 1


    print(f"    → {len(results)} chunks from {len(reader.pages)} pages")
    return results
        
        
def read_csv(filepath):
    """
    Reads a CSV file row by row.

    What:  Reads each row → converts to readable text sentence → stores with row metadata
    Why:   CSV rows are structured data — each row = one chunk
           We convert row to natural language so LLM can read it easily
    Returns: list of (chunk_text, metadata_dict) tuples

    Example:
        Row: course_name=Machine Learning, price=4999, duration=3 months
        Text: "course_name: Machine Learning | price: 4999 | duration: 3 months"
    """
    filename = os.path.basename(filepath)
    print(f"  Reading CSV: {filename}")

    # pandas reads CSV into a DataFrame (like a table)
    df = pd.read_csv(filepath)

    results = []

    for row_index, row in df.iterrows():
        # Convert each row to a readable text string
        # row.items() gives (column_name, value) pairs
        # We join them as "column: value | column: value | ..."
        row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])

        metadata = {
            "source": filename,
            "file_type": "csv",
            "row": int(row_index) + 1,     # row number (1-based)
            "chunk_index": int(row_index),
            "category": "data"
        }

        results.append((row_text, metadata))

    print(f"    → {len(results)} chunks (one per row)")
    return results


def read_excel(filepath):
    """
    Reads an Excel file sheet by sheet, row by row.

    What:  Reads each sheet → each row → converts to text → stores with metadata
    Why:   Excel can have multiple sheets — we track sheet name in metadata
    Returns: list of (chunk_text, metadata_dict) tuples
    """
    filename = os.path.basename(filepath)
    print(f"  Reading Excel: {filename}")

    # pd.ExcelFile lets us access all sheets
    excel_file = pd.ExcelFile(filepath)
    results = []
    chunk_index = 0

    # Loop through each sheet
    for sheet_name in excel_file.sheet_names:

        # Read this sheet into a DataFrame
        df = pd.read_excel(filepath, sheet_name=sheet_name)

        for row_index, row in df.iterrows():

            # Convert row to readable text (same as CSV)
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])

            metadata = {
                "source": filename,
                "file_type": "excel",
                "sheet": sheet_name,           # which sheet
                "row": int(row_index) + 1,
                "chunk_index": chunk_index,
                "category": "data"
            }

            results.append((row_text, metadata))
            chunk_index += 1

    print(f"    → {len(results)} chunks from {len(excel_file.sheet_names)} sheets")
    return results

def read_docx(filepath):
    """
    Reads a Word document paragraph by paragraph.

    What:  Reads each paragraph → chunks if too long → stores with metadata
    Why:   Word docs are structured by paragraphs — natural chunking boundary
    Returns: list of (chunk_text, metadata_dict) tuples
    """
    filename = os.path.basename(filepath)
    print(f"  Reading DOCX: {filename}")

    doc = Document(filepath)
    results = []
    chunk_index = 0

    for para_index, paragraph in enumerate(doc.paragraphs):

        # Get paragraph text
        para_text = paragraph.text.strip()

        # Skip empty paragraphs
        if not para_text:
            continue

        # If paragraph is small enough — keep as one chunk
        if len(para_text) <= CHUNK_SIZE:
            metadata = {
                "source": filename,
                "file_type": "docx",
                "paragraph": para_index + 1,
                "chunk_index": chunk_index,
                "category": "document"
            }
            results.append((para_text, metadata))
            chunk_index += 1

        else:
            # Paragraph too long — chunk further
            sub_chunks = chunk_text(para_text)
            for sub_index, sub_chunk in enumerate(sub_chunks):
                metadata = {
                    "source": filename,
                    "file_type": "docx",
                    "paragraph": para_index + 1,
                    "sub_chunk": sub_index,
                    "chunk_index": chunk_index,
                    "category": "document"
                }
                results.append((sub_chunk, metadata))
                chunk_index += 1

    print(f"    → {len(results)} chunks from {len(doc.paragraphs)} paragraphs")
    return results


def read_file(filepath):
    """
    Router function - reads any file type and dispatches to the correct reader
    
    What: Determines file type by extension and calls appropriate reader
    Why: Allows Load_all_docs() to process mixed file types seamlessly
    Returns: list of (chunk_text, metadata_dict) tuples
    """
    # Get file extension
    _, file_extension = os.path.splitext(filepath)
    file_extension = file_extension.lower()
    
    # Skip hidden files and common non-document files
    filename = os.path.basename(filepath)
    if filename.startswith('.'):
        return []
    
    # Route to appropriate reader
    if file_extension == '.txt':
        return read_txt(filepath)
    elif file_extension == '.pdf':
        return read_pdf(filepath)
    elif file_extension == '.csv':
        return read_csv(filepath)
    elif file_extension in ['.xlsx', '.xls']:
        return read_excel(filepath)
    elif file_extension == '.docx':
        return read_docx(filepath)
    else:
        print(f"  ⚠️  Unsupported file type: {file_extension}")
        return []