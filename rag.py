# RAG with Chroma DB
from groq import Groq
from dotenv import load_dotenv
import os
import chromadb
from chromadb.utils import embedding_functions

#--------------------For reading libraries----------------------
import pandas as pd                 #For reading CSV and excel files
from pypdf import PdfReader         #For reading PDF files
from docx import Document           #For word files


#-----------------------for smart chunking--------------------
import re                           #for regular expression, to split text into paragraphs



#---------------------Full project pipeline------------------------
''' 
FULL PROJECT PIPELINE:
1. get collection: create or get chroma db collection
2. load and store knowledge:  load knowledge from knowledge.txt file and separate into chunks and sotre in db
3. user ask a question
4. retrieve relavant chunks: based on the question, retrieve relavant chunks from db
5. build prompt: combime the chunks and the question to build a prompt for llm
6. generate answer: send the prompt to llm and get answer
'''


#---------LOAD ENV-------------
load_dotenv()

#-----------CONSTANTS----------------
KNOWLEDGE_FILE = "knowledge.txt"
CHROMA_DB_PATH =  "./chroma_db"
#-------------Chunking settings---------------
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

#------------COLLECTION-----------------
#This is the name of collection, same as the name of Table in SQL
DOCUMENTS_FOLDER = "./documents"
COLLECTION_NAME = "multi_file_rag"
MODEL = "llama-3.3-70b-versatile"

#---------------------GROQ CLIENT----------------------
client = Groq(api_key = os.environ.get("GROQ_API_KEY"))


#---------------------EMBEDDING FUNCION----------------------
# How:   We use sentence-transformers — a free local model
#        "all-MiniLM-L6-v2" is small, fast, and works well for this use case
#        It converts any text into a vector of 384 numbers
'''embedding function will convert the text into vector'''
embedding_fn = embedding_functions.SentenceTransformerEmbeddingFunction(model_name = "all-MiniLM-L6-v2")


#--------------------SMART TEXT CHUNKING ----------------------------
def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    '''
    splits the contionious text into overlapping chunks
    tries to split the sentance from (.) then words
    '''
    #step 1 :- clean the text
    #replace the multiple spaces/newlines with single space
    text = re.sub(r'\s+','',text).strip()
    
    #if text is smaller than chunk_size - return as single chunk
    if len(text) <= CHUNK_SIZE:
        return[text]
    
    chunks = []
    start = 0
    #here first find the starting and ending index in the chunk and then 
    #extract the chunk on the basis of starting and ending index
    while start < len(text):
        end = start+chunk_size
        #if we haven't reached the end yet
        if end < len(text):
            #trying to find the last sentance ending (.)
            #to prevent the text from cutting from middle
            last_period = text.rfind('.',start,end) #finding last occurence of (.) from start to end
            
            if last_period != -1  and last_period > start + (chunk_size // 2):
                #we got the text in second half
                # cut from there instead of exact chunk size
                end = last_period + 1
            else:
                last_space = text.rfind(' ',start,end)
                if last_space != -1:
                    end = last_space
        #extracting chunk
        chunk = text[start:end].strip()
        
        #only add non empty chunks
        if chunk :
            chunks.append(chunk)            
        
        #moving start bid backward for overlapping
        start = end - overlap
        
        if start >= end:
            start = end
            
            
    return chunks
                


#--------------------FILE READER----------------------------
def read_txt(filepath):
    '''Reads a plain text file and chunks it smartly
    1.WHAT :- read the .txt file and chunks it smartly
    2.RETURNS :- list of (chunks_text, metadata_dict)
    '''
    filename = os.path.basename(filepath)
    print(f"Reading text:{filename}")
    
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    
    #first trying to split by double newline (paragraph)
    paragraphs = [p.strip() for p in content .split("\n\n") if p.strip()]
    
    results = []
    chunk_index = 0
    
    for para_index, paragraph in enumerate(paragraphs):
        #if paragraph is small - keep in one chunk
        if len(paragraph) <= CHUNK_SIZE:
            metadata = {
                "source" : filename,
                "file_type": "txt",
                "paragraph": para_index,
                "chunk_index": chunk_index,
                "category" : "text"
            }
            results.append((paragraph,metadata))
            chunk_index +- 1
    
    print(f"   -> {len(results)}chunks created")
    return results

#----------------------read pdf----------------------------
def read_pdf(filepath):
    filename = os.path.basename(filepath)
    print(f"reading pdf: {filename}")
    
    reader = PdfReader(filepath)
    results = []
    chunk_index = 0
    
    for page_num, page in enumerate(reader.pages):
        #extracting text from the page
        page_text = page.extract_text()
        
        if not page_text or not page_text.strip():
            continue
        
        chunks = chunk_text(page_text)
        for sub_index, chunk in enumerate(chunks):
            metadata = {
                "source": filename,
                "file_type": "pdf",
                "page": page_num + 1,          # page numbers start at 1
                "sub_chunk": sub_index,
                "chunk_index": chunk_index,
                "category": "document"
            }
            results.append((chunk, metadata))
            chunk_index += 1


    print(f"    → {len(results)} chunks from {len(reader.pages)} pages")
    return results
        
        
def read_csv(filepath):
    """
    Reads a CSV file row by row.

    What:  Reads each row → converts to readable text sentence → stores with row metadata
    Why:   CSV rows are structured data — each row = one chunk
           We convert row to natural language so LLM can read it easily
    Returns: list of (chunk_text, metadata_dict) tuples

    Example:
        Row: course_name=Machine Learning, price=4999, duration=3 months
        Text: "course_name: Machine Learning | price: 4999 | duration: 3 months"
    """
    filename = os.path.basename(filepath)
    print(f"  Reading CSV: {filename}")

    # pandas reads CSV into a DataFrame (like a table)
    df = pd.read_csv(filepath)

    results = []

    for row_index, row in df.iterrows():
        # Convert each row to a readable text string
        # row.items() gives (column_name, value) pairs
        # We join them as "column: value | column: value | ..."
        row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])

        metadata = {
            "source": filename,
            "file_type": "csv",
            "row": int(row_index) + 1,     # row number (1-based)
            "chunk_index": int(row_index),
            "category": "data"
        }

        results.append((row_text, metadata))

    print(f"    → {len(results)} chunks (one per row)")
    return results


def read_excel(filepath):
    """
    Reads an Excel file sheet by sheet, row by row.

    What:  Reads each sheet → each row → converts to text → stores with metadata
    Why:   Excel can have multiple sheets — we track sheet name in metadata
    Returns: list of (chunk_text, metadata_dict) tuples
    """
    filename = os.path.basename(filepath)
    print(f"  Reading Excel: {filename}")

    # pd.ExcelFile lets us access all sheets
    excel_file = pd.ExcelFile(filepath)
    results = []
    chunk_index = 0

    # Loop through each sheet
    for sheet_name in excel_file.sheet_names:

        # Read this sheet into a DataFrame
        df = pd.read_excel(filepath, sheet_name=sheet_name)

        for row_index, row in df.iterrows():

            # Convert row to readable text (same as CSV)
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])

            metadata = {
                "source": filename,
                "file_type": "excel",
                "sheet": sheet_name,           # which sheet
                "row": int(row_index) + 1,
                "chunk_index": chunk_index,
                "category": "data"
            }

            results.append((row_text, metadata))
            chunk_index += 1

    print(f"    → {len(results)} chunks from {len(excel_file.sheet_names)} sheets")
    return results


def read_docx(filepath):
    """
    Reads a Word document paragraph by paragraph.

    What:  Reads each paragraph → chunks if too long → stores with metadata
    Why:   Word docs are structured by paragraphs — natural chunking boundary
    Returns: list of (chunk_text, metadata_dict) tuples
    """
    filename = os.path.basename(filepath)
    print(f"  Reading DOCX: {filename}")

    doc = Document(filepath)
    results = []
    chunk_index = 0

    for para_index, paragraph in enumerate(doc.paragraphs):

        # Get paragraph text
        para_text = paragraph.text.strip()

        # Skip empty paragraphs
        if not para_text:
            continue

        # If paragraph is small enough — keep as one chunk
        if len(para_text) <= CHUNK_SIZE:
            metadata = {
                "source": filename,
                "file_type": "docx",
                "paragraph": para_index + 1,
                "chunk_index": chunk_index,
                "category": "document"
            }
            results.append((para_text, metadata))
            chunk_index += 1

        else:
            # Paragraph too long — chunk further
            sub_chunks = chunk_text(para_text)
            for sub_index, sub_chunk in enumerate(sub_chunks):
                metadata = {
                    "source": filename,
                    "file_type": "docx",
                    "paragraph": para_index + 1,
                    "sub_chunk": sub_index,
                    "chunk_index": chunk_index,
                    "category": "document"
                }
                results.append((sub_chunk, metadata))
                chunk_index += 1

    print(f"    → {len(results)} chunks from {len(doc.paragraphs)} paragraphs")
    return results




#---------------------READ ROUTER --------------------------
def read_file(filepath):
    '''route each file to the correct reader'''
    extension = os.path.splitext(filepath)[1].lower()
    
    #map extenstion to render
    readers = {
        ".txt":read_txt,
        ".pdf":read_pdf,
        ".csv":read_csv,
        ".xlsx":read_excel,
        ".docx":read_docx
    }
    
    if extension in readers:
        return readers[extension](filepath)
    else:
        print(f" skipping unsupported file: {filepath}")
        return[]


#---------------------CHROMA DB SETUP---------------------
def get_collection():
    #create a chroma db client
    chroma_client = chromadb.PersistentClient(path = CHROMA_DB_PATH)
    
    #check if collection exists, if not create it
    collection = chroma_client.get_or_create_collection(
        name = COLLECTION_NAME,
        embedding_function = embedding_fn
    )
    
    return collection


#--------------------LOAD AND STORE KNOWLEDGE---------------------
def Load_all_docs(collection):
    '''
    1. it is one function to load everything
    2. loops through the files -> routes to the correct reader 
    3. read the content -> split into chunks -> store in chroma db
    '''
    #load the knowledge.txt file, separate into chunks and store in chroma db
    if collection.count() > 0:
        print(f"collection already has {collection.count()} chunks, skipping loading")
        return
    print(f"Loading all documents into ChromaDB...\n")
    
    if not os.path.exists(DOCUMENTS_FOLDER):
        print(f"Documents folder '{DOCUMENTS_FOLDER}' not found. Please create it and add  your documents.")
        return
    
    #getting all files of documents folder
    all_files = os.listdir(DOCUMENTS_FOLDER)
    
    all_chunks=[]
    all_metadata=[]
    all_ids=[]
    
    #global chunk counter for uniqueid
    global_id = 0
    
    #process each file
    for filename in all_files:
        filepath = os.path.join(DOCUMENTS_FOLDER, filename)
        #read the file -> get list of (text,metadata) tuples
        file_chunks = read_file(filepath)
        for text,metadata in file_chunks:
            all_chunks.append(text)
            all_metadata.append(metadata)
            all_ids.append(f"chunk_{global_id}")
            global_id += 1
            
    if not all_chunks:
        print(f"No chunks found for this question in this folder")
        return
    
    print(f"storing {len(all_chunks)} total chunks in ChromaDB....")
    
    '''AS CHROMA DB IS HAVING LIMIT PER BATCH..Storing in batch of 100'''
    batch_size = 100
    for i in range(0,len(all_chunks), batch_size):
        batch_chunks = all_chunks[i : i+batch_size]
        batch_metadata = all_metadata[i : i+batch_size]
        batch_ids = all_ids[i : i+batch_size]
        
        collection.add(
            documents = batch_chunks,
            metadatas = batch_metadata,
            ids = batch_ids
        )
        print(f"  Stored batch {i//batch_size + 1} ({len(batch_chunks)} chunks)")
        
    print(f"\nDone! {len(all_chunks)} chunks stored in ChromaDB.\n")
    
    
#--------------------RETRIEVE RELAVANT CHUNKS----------------------
def retrieve(question, collection, top_k = 3):
    #retrieve relavant chunks from chroma db based on the question\
        
    '''the chromadb convert the question to the vector using the embeding function and then 
    compare the vector and return the most similar chunks'''
    results = collection.query(
        query_texts = [question],
        n_results = top_k
    )
    
    chunks = results["documents"][0]
    metadatas = results["metadatas"][0]
    return chunks,metadatas


#---------------------build prompt with context-----------------------
def build_prompt(question, chunks, metadatas):
    # Build prompt with context and souce information
    # Each chunk is labelled with its source file so that llm can site it.
    
    context_parts = []
    
    for chunk,metadata in zip(chunks, metadatas):  #zip, pairs the elements of both lists
        source = metadata.get("source", "unknown")
        file_type = metadata.get("file_type", "unknown")
        
        #add extra location info based on file type
        if file_type == "pdf":
            location = f"Page {metadata.get('page', '?')}"
        elif file_type in ["csv", "excel"]:
            location = f"Row {metadata.get('row', '?')}"
        elif file_type == "docx":
            location = f"Paragraph {metadata.get('paragraph', '?')}"
        else:
            location = f"Paragraph {metadata.get('paragraph', '?')}"
            
        # Format: [Source: policy.pdf, Page 2]
        source_label = f"[Source: {source}, {location}]"

        context_parts.append(f"{source_label}\n{chunk}")
 
        
    context = "\n\n".join(context_parts)
    
    prompt = f"""You are a helpful assistant for TechLearn India.
    Use the following context to answer the question.
    Always mention which source/file the information came from.
    If the answer is not in the context say "I don't have that information."
    Context: {context}

    Customer Question: {question}

    Answer: """
    return prompt
    
    
#---------------------generate answer----------------------
def ask(question,collection):
    ''' Retrieve the relavant chunks from chroma db,
    build prompt and generate answer
    give it to  llm
    return answer'''
    print(f"Getting the relavant chunks from the ChromaDB for the question")
    chunks,metadatas = retrieve(question, collection)
    
    for i, (chunk,meta) in enumerate(zip(chunks,metadatas)):
        print(f"  [{i+1}] {meta['source']} | {chunk[:50]}...")
        
        
    prompt = build_prompt(question, chunks, metadatas)
    
    #sending to llm
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": "You are a helpful assistant for TechLearn India. Always cite your sources."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        temperature=0.3,
        max_tokens=512
    )
    answer = response.choices[0].message.content
    return answer


def main():
    print("=" * 55)
    print("   📚 TechLearn RAG with ChromaDB")
    print("=" * 55 + "\n")
    #get collections
    collection = get_collection()
    #load store knowledge
    Load_all_docs(collection)
    
    print(f"✅ Ready! {collection.count()} chunks in database.\n")
    print("Type 'quit' to exit | 'stats' to see DB info\n")
    
    #ask question
    while True:
        question = input("You: ").strip()

        if not question:
            continue

        if question.lower() == "quit":
            print("Goodbye!")
            break

        elif question.lower() == "stats":
            print(f"\n📊 ChromaDB Stats:")
            print(f"   Total chunks: {collection.count()}")
            print(f"   DB location: {CHROMA_DB_PATH}\n")
            continue

        answer = ask(question, collection)
        print(f"\nAssistant: {answer}\n")
        print("-" * 50)

        
        
if __name__ == "__main__":
    main()